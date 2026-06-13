from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "documents" / "japanese_ocr_sample.docx"
TEXT_PDF = ROOT / "documents" / "japanese_ocr_text.pdf"
IMAGE_PDF = ROOT / "documents" / "japanese_ocr_image.pdf"
PHOTO = ROOT / "documents" / "japanese_ocr_photo.png"
IMAGE_PREVIEW = ROOT / "results" / "japanese_ocr_image_source.png"
FONT = "Yu Gothic"
BLUE = "2E5D7B"
LIGHT_BLUE = "E8F1F6"
LIGHT_GRAY = "F2F4F7"


def set_run_font(run, size=11, bold=False, color="222222"):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_text(paragraph, text, size=11, bold=False, color="222222"):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return run


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.15

for style_name, size, before, after in (
    ("Heading 1", 15, 10, 5),
    ("Heading 2", 12, 7, 3),
):
    style = doc.styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(BLUE)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_text(header, "OCR評価用サンプル / 2026-06-13", size=8.5, color="6B7280")

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(2)
add_text(title, "日本語 OCR・文書抽出 評価サンプル", size=20, bold=True, color=BLUE)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(8)
add_text(
    subtitle,
    "画像PDFとテキストPDFで、文字認識・表構造・記号の再現性を比較するための原稿",
    size=9.5,
    color="4B5563",
)

note = doc.add_table(rows=1, cols=1)
set_table_geometry(note, [9840])
shade_cell(note.cell(0, 0), LIGHT_BLUE)
p = note.cell(0, 0).paragraphs[0]
p.paragraph_format.space_after = Pt(0)
add_text(
    p,
    "確認ポイント: 漢字とかな、全角・半角、似た字形、住所、日付、金額、表の行列、句読点。",
    size=9.5,
    bold=True,
    color=BLUE,
)

doc.add_heading("1. 基本文字と文章", level=1)
p = doc.add_paragraph()
add_text(
    p,
    "春の朝、東京都千代田区丸の内一丁目では、桜の花びらが静かに舞っていました。"
    "受付番号は「京A-0427」、申請日は2026年6月13日、締切は同年6月30日です。",
)
p = doc.add_paragraph()
add_text(p, "ひらがな: あいうえお・がぎぐげご・ぱぴぷぺぽ　")
add_text(p, "カタカナ: アイウエオ・ヴァイオリン・コンピューター")

doc.add_heading("2. 誤認識しやすい文字列", level=1)
tests = [
    ("英数字", "O0Q / I1l / S5 / B8 / Z2 / rn-m / cl-d"),
    ("全角半角", "ＡＢＣ１２３ / ABC123 / ﾊﾝｶｸｶﾅ / 半角カナ"),
    ("記号", "「」『』（）［］【】・…―〜／￥〒 ※ ○ × △ □"),
    ("コード", "INV-2026-00081 / JP13-0001-0427 / 03-1234-5678"),
]
for label, value in tests:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    add_text(p, f"{label}: ", bold=True, color=BLUE)
    add_text(p, value)

doc.add_heading("3. 帳票形式の表", level=1)
rows = [
    ["品番", "品名", "数量", "単価（税込）", "備考"],
    ["A-001", "抹茶入り緑茶 500ml", "12", "¥1,280", "賞味期限 2027/03"],
    ["B-017", "国産りんご（ふじ）", "8", "¥3,456", "青森県産"],
    ["C-105", "精密ねじ M3×12", "250", "¥9.80", "0/O の判別"],
    ["D-808", "USB-C ケーブル 1.5m", "3", "¥2,200", "I/1/l の判別"],
]
table = doc.add_table(rows=len(rows), cols=5)
table.style = "Table Grid"
set_table_geometry(table, [1250, 3300, 900, 1700, 2690])
for r, values in enumerate(rows):
    for c, value in enumerate(values):
        cell = table.cell(r, c)
        if r == 0:
            shade_cell(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c in (0, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        add_text(p, value, size=9, bold=r == 0, color=BLUE if r == 0 else "222222")

doc.add_heading("4. 連絡先とチェック項目", level=1)
p = doc.add_paragraph()
add_text(p, "株式会社 青空物流　総務部 文書管理課\n", bold=True)
add_text(p, "〒100-0005 東京都千代田区丸の内1-2-3 青空ビル7階\n")
add_text(p, "電話: 03-1234-5678　メール: ocr-test@example.jp")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(3)
add_text(p, "☑ 受領済み　☐ 要確認　☒ 不備あり　")
add_text(p, "小さい文字: これは9ポイントの注記です。", size=9, color="777777")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(footer, "機密情報を含まない架空データです。", size=8, color="777777")

doc.add_page_break()
photo_title = doc.add_paragraph()
photo_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
photo_title.paragraph_format.space_after = Pt(6)
add_text(photo_title, "5. スキャン風帳票画像", size=15, bold=True, color=BLUE)
photo_paragraph = doc.add_paragraph()
photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
photo_paragraph.paragraph_format.space_after = Pt(0)
photo_paragraph.add_run().add_picture(str(PHOTO), width=Inches(5.8))

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)


pdfmetrics.registerFont(UnicodeCIDFont("HeiseiKakuGo-W5"))
pdfmetrics.registerFont(UnicodeCIDFont("HeiseiMin-W3"))


def draw_text_pdf():
    page_width, page_height = A4
    pdf = canvas.Canvas(str(TEXT_PDF), pagesize=A4)
    pdf.setTitle("Japanese OCR Text Sample")

    def text(x, y, value, size=10, font="HeiseiKakuGo-W5", color=colors.HexColor("#222222")):
        pdf.setFont(font, size)
        pdf.setFillColor(color)
        pdf.drawString(x, y, value)

    text(42, page_height - 48, "日本語 OCR・文書抽出 評価サンプル", 18, color=colors.HexColor("#2E5D7B"))
    text(
        42,
        page_height - 66,
        "漢字・かな・英数字・記号・住所・日付・金額・表構造の確認用",
        9,
        color=colors.HexColor("#555555"),
    )

    y = page_height - 92
    text(42, y, "1. 基本文字と文章", 12, color=colors.HexColor("#2E5D7B"))
    y -= 18
    text(42, y, "春の朝、東京都千代田区丸の内一丁目では、桜の花びらが静かに舞っていました。", 9.5)
    y -= 15
    text(42, y, "受付番号は「京A-0427」、申請日は2026年6月13日、締切は2026年6月30日です。", 9.5)
    y -= 15
    text(42, y, "ひらがな: あいうえお・がぎぐげご・ぱぴぷぺぽ", 9.5)
    y -= 15
    text(42, y, "カタカナ: アイウエオ・ヴァイオリン・コンピューター", 9.5)

    y -= 24
    text(42, y, "2. 誤認識しやすい文字列", 12, color=colors.HexColor("#2E5D7B"))
    for value in (
        "英数字: O0Q / I1l / S5 / B8 / Z2 / rn-m / cl-d",
        "全角半角: ＡＢＣ１２３ / ABC123 / ﾊﾝｶｸｶﾅ / 半角カナ",
        "記号: 「」『』（）［］【】・…―〜／￥〒 ※ ○ × △ □",
        "コード: INV-2026-00081 / JP13-0001-0427 / 03-1234-5678",
    ):
        y -= 16
        text(48, y, value, 9.5)

    y -= 25
    text(42, y, "3. 帳票形式の表", 12, color=colors.HexColor("#2E5D7B"))
    y -= 15
    x_positions = [42, 112, 286, 340, 425, 553]
    table_rows = [
        ["品番", "品名", "数量", "単価（税込）", "備考"],
        ["A-001", "抹茶入り緑茶 500ml", "12", "¥1,280", "賞味期限 2027/03"],
        ["B-017", "国産りんご（ふじ）", "8", "¥3,456", "青森県産"],
        ["C-105", "精密ねじ M3×12", "250", "¥9.80", "0/O の判別"],
        ["D-808", "USB-C ケーブル 1.5m", "3", "¥2,200", "I/1/l の判別"],
    ]
    row_height = 24
    for row_index, row in enumerate(table_rows):
        top = y - row_index * row_height
        if row_index == 0:
            pdf.setFillColor(colors.HexColor("#E8F1F6"))
            pdf.rect(x_positions[0], top - row_height, x_positions[-1] - x_positions[0], row_height, fill=1, stroke=0)
        for x in x_positions:
            pdf.setStrokeColor(colors.HexColor("#A8B6C0"))
            pdf.line(x, top, x, top - row_height)
        pdf.line(x_positions[0], top, x_positions[-1], top)
        for column, value in enumerate(row):
            text(
                x_positions[column] + 4,
                top - 16,
                value,
                7.8 if row_index else 8.2,
                color=colors.HexColor("#2E5D7B") if row_index == 0 else colors.HexColor("#222222"),
            )
    bottom = y - len(table_rows) * row_height
    pdf.line(x_positions[0], bottom, x_positions[-1], bottom)

    y = bottom - 28
    text(42, y, "4. 連絡先とチェック項目", 12, color=colors.HexColor("#2E5D7B"))
    y -= 18
    text(42, y, "株式会社 青空物流　総務部 文書管理課", 9.5)
    y -= 15
    text(42, y, "〒100-0005 東京都千代田区丸の内1-2-3 青空ビル7階", 9.5)
    y -= 15
    text(42, y, "電話: 03-1234-5678　メール: ocr-test@example.jp", 9.5)
    y -= 20
    text(42, y, "■ 受領済み　□ 要確認　× 不備あり", 9.5)
    text(330, y, "小さい薄色文字: これは判読しにくい注記です。", 7, color=colors.HexColor("#999999"))

    text(42, 30, "機密情報を含まない架空データです。", 7.5, color=colors.HexColor("#777777"))
    pdf.save()


def draw_image_pdf():
    scale = 3
    width, height = (int(A4[0] * scale), int(A4[1] * scale))
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    regular_path = r"C:\Windows\Fonts\NotoSansJP-VF.ttf"
    bold_path = r"C:\Windows\Fonts\NotoSansJP-VF.ttf"

    def font(size):
        return ImageFont.truetype(regular_path, int(size * scale))

    def put(x, y, value, size=10, color="#222222"):
        draw.text((int(x * scale), int(y * scale)), value, font=font(size), fill=color)

    put(42, 34, "日本語 OCR・文書抽出 評価サンプル", 18, "#2E5D7B")
    put(42, 57, "漢字・かな・英数字・記号・住所・日付・金額・表構造の確認用", 9, "#555555")
    y = 90
    put(42, y, "1. 基本文字と文章", 12, "#2E5D7B")
    y += 22
    put(42, y, "春の朝、東京都千代田区丸の内一丁目では、桜の花びらが静かに舞っていました。", 9.5)
    y += 18
    put(42, y, "受付番号は「京A-0427」、申請日は2026年6月13日、締切は2026年6月30日です。", 9.5)
    y += 18
    put(42, y, "ひらがな: あいうえお・がぎぐげご・ぱぴぷぺぽ", 9.5)
    y += 18
    put(42, y, "カタカナ: アイウエオ・ヴァイオリン・コンピューター", 9.5)
    y += 28
    put(42, y, "2. 誤認識しやすい文字列", 12, "#2E5D7B")
    for value in (
        "英数字: O0Q / I1l / S5 / B8 / Z2 / rn-m / cl-d",
        "全角半角: ＡＢＣ１２３ / ABC123 / ﾊﾝｶｸｶﾅ / 半角カナ",
        "記号: 「」『』（）［］【】・…―〜／￥〒 ※ ○ × △ □",
        "コード: INV-2026-00081 / JP13-0001-0427 / 03-1234-5678",
    ):
        y += 19
        put(48, y, value, 9.5)

    y += 30
    put(42, y, "3. 帳票形式の表", 12, "#2E5D7B")
    y += 23
    x_positions = [42, 112, 286, 340, 425, 553]
    table_rows = [
        ["品番", "品名", "数量", "単価（税込）", "備考"],
        ["A-001", "抹茶入り緑茶 500ml", "12", "¥1,280", "賞味期限 2027/03"],
        ["B-017", "国産りんご（ふじ）", "8", "¥3,456", "青森県産"],
        ["C-105", "精密ねじ M3×12", "250", "¥9.80", "0/O の判別"],
        ["D-808", "USB-C ケーブル 1.5m", "3", "¥2,200", "I/1/l の判別"],
    ]
    row_height = 26
    for row_index, row in enumerate(table_rows):
        top = y + row_index * row_height
        if row_index == 0:
            draw.rectangle(
                (int(x_positions[0] * scale), int(top * scale), int(x_positions[-1] * scale), int((top + row_height) * scale)),
                fill="#E8F1F6",
            )
        for x in x_positions:
            draw.line((int(x * scale), int(top * scale), int(x * scale), int((top + row_height) * scale)), fill="#A8B6C0", width=2)
        draw.line((int(x_positions[0] * scale), int(top * scale), int(x_positions[-1] * scale), int(top * scale)), fill="#A8B6C0", width=2)
        for column, value in enumerate(row):
            put(x_positions[column] + 4, top + 7, value, 7.8 if row_index else 8.2, "#2E5D7B" if row_index == 0 else "#222222")
    bottom = y + len(table_rows) * row_height
    draw.line((int(x_positions[0] * scale), int(bottom * scale), int(x_positions[-1] * scale), int(bottom * scale)), fill="#A8B6C0", width=2)

    y = bottom + 27
    put(42, y, "4. 連絡先とチェック項目", 12, "#2E5D7B")
    y += 23
    put(42, y, "株式会社 青空物流　総務部 文書管理課", 9.5)
    y += 18
    put(42, y, "〒100-0005 東京都千代田区丸の内1-2-3 青空ビル7階", 9.5)
    y += 18
    put(42, y, "電話: 03-1234-5678　メール: ocr-test@example.jp", 9.5)
    y += 24
    put(42, y, "■ 受領済み　□ 要確認　× 不備あり", 9.5)
    put(330, y + 2, "小さい薄色文字: これは判読しにくい注記です。", 7, "#999999")
    put(42, 805, "機密情報を含まない架空データです。", 7.5, "#777777")

    IMAGE_PREVIEW.parent.mkdir(parents=True, exist_ok=True)
    image.save(IMAGE_PREVIEW)
    photo = Image.open(PHOTO).convert("RGB")
    photo_page = Image.new("RGB", (width, height), "white")
    photo.thumbnail((width - 180, height - 180), Image.Resampling.LANCZOS)
    photo_page.paste(photo, ((width - photo.width) // 2, (height - photo.height) // 2))
    image.save(IMAGE_PDF, "PDF", resolution=216, save_all=True, append_images=[photo_page])


draw_text_pdf()
draw_image_pdf()
print(TEXT_PDF)
print(IMAGE_PDF)
